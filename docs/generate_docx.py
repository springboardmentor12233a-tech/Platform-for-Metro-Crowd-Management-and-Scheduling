from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page 1: Title Page
title_paragraphs = [
    ("\n\n\n\n\n\n", 12),
    ("MetroFlow: AI Platform for Metro Crowd Management and Scheduling", 24),
    ("\n\n\n\n", 12),
    ("Milestone 1 Report", 18),
    ("\n\n\n\n", 12),
    ("Prepared by: Ankita Jana", 14),
    ("\n\n", 12),
    ("Date: July 2026", 12)
]

for text, size in title_paragraphs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if size > 12:
        run.bold = True
    run.font.size = Pt(size)

doc.add_page_break()

# Page 2: Objective & Project Overview
doc.add_heading('2. Objective & Project Overview', level=1)

doc.add_heading('2.1 Project Objective', level=2)
p = doc.add_paragraph("Build an AI-powered metro crowd management and scheduling platform that helps metro authorities monitor passenger flow, predict crowd density, and optimize train scheduling in real time. ")
p = doc.add_paragraph("The system analyzes passenger traffic patterns, station congestion, train occupancy, ticketing data, sensor data, and peak-hour demand to improve metro operations and reduce overcrowding.")
p = doc.add_paragraph("The platform is designed to support smart transportation systems by integrating AI analytics, scheduling automation, crowd prediction, and operational monitoring into a centralized application.")
p = doc.add_paragraph("This solution is intended for metro rail systems, smart city transportation, urban mobility management, and public transit optimization.")

doc.add_heading('2.2 Expected Outcomes', level=2)
outcomes = [
    "Developed and deployed an AI-powered metro crowd management platform.",
    "Implemented real-time passenger density and congestion monitoring systems.",
    "Built AI-based crowd prediction and passenger demand forecasting models.",
    "Developed train scheduling and frequency optimization workflows.",
    "Implemented alert and emergency notification systems.",
    "Built analytics dashboards for passenger traffic and operational monitoring.",
    "Developed scalable Python-based backend services for real-time transportation analytics.",
    "Deployed the platform using Docker and cloud deployment platforms such as AWS or Azure."
]
for out in outcomes:
    doc.add_paragraph(out, style='List Bullet')

doc.add_heading('2.3 Milestone 1 Goals', level=2)
p = doc.add_paragraph()
run = p.add_run("Week 1 & 2 — Project Initialization, Design Process & Core Setup")
run.bold = True
goals = [
    "Define project objectives and transportation workflows.",
    "Design system architecture and database schema.",
    "Create UI wireframes and operational workflow planning.",
    "Setup frontend and backend environments.",
    "Implement authentication and role-based access system.",
    "Build crowd monitoring dashboard.",
    "Develop congestion tracking features."
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run("Milestone Outcomes:")
run.bold = True
m_outcomes = [
    "Understand smart transportation and metro management workflows.",
    "Learn system architecture and database design concepts.",
    "Build frontend and backend project initialization.",
    "Working authentication and crowd monitoring system."
]
for mo in m_outcomes:
    doc.add_paragraph(mo, style='List Bullet')

doc.add_page_break()

# Page 3
doc.add_heading('3. Design Process & Core Setup', level=1)

doc.add_heading('3.1 Defining Transportation Workflows', level=2)
doc.add_paragraph("The first step in our achievement process involved outlining the core operational workflows of a standard metro transportation system. This included:")
doc.add_paragraph("Passenger Flow Tracking: From entry at ticketing gates to exit, analyzing dwell time.", style='List Bullet')
doc.add_paragraph("Train Scheduling: Timetables, frequency during peak vs. non-peak hours, and delay management.", style='List Bullet')
doc.add_paragraph("Congestion Alerts: Defining thresholds for crowd density at platforms and inside train cars to trigger alerts.", style='List Bullet')

doc.add_heading('3.2 System Architecture Design', level=2)
doc.add_paragraph("A scalable, microservices-oriented architecture was designed to handle real-time data ingestion and processing.")
doc.add_paragraph("Frontend Layer: A React/Vite based Single Page Application (SPA) providing an interactive dashboard for operators.", style='List Bullet')
doc.add_paragraph("Backend Layer: A Python (FastAPI/Django) backend to serve RESTful APIs, manage business logic, and handle ML inference for predictive models.", style='List Bullet')
doc.add_paragraph("Data Layer: PostgreSQL for structured relational data (users, schedules, station info) and MongoDB/Redis for high-velocity telemetry data (sensor readings, crowd density).", style='List Bullet')
doc.add_paragraph("AI/ML Layer: Dedicated Python services leveraging PyTorch/Scikit-Learn to forecast ridership and detect anomalies.", style='List Bullet')

doc.add_heading('3.3 Database Schema Design', level=2)
doc.add_paragraph("We established a normalized relational schema for the core application:")
doc.add_paragraph("Users: ID, Name, Role (Admin, Operator, Analyst), Credentials.", style='List Bullet')
doc.add_paragraph("Stations: StationID, Name, Line, Zone, Capacity.", style='List Bullet')
doc.add_paragraph("Trains: TrainID, Capacity, CurrentLine, Status.", style='List Bullet')
doc.add_paragraph("Schedules: ScheduleID, TrainID, StationID, ArrivalTime, DepartureTime.", style='List Bullet')
doc.add_paragraph("Telemetry: LogID, Timestamp, Location, DensityScore.", style='List Bullet')

doc.add_heading('3.4 UI Wireframing', level=2)
doc.add_paragraph("Initial wireframes were created to visualize the Analytics Dashboard, Schedule Manager, and Alert Center. The design focuses on high visibility of critical metrics (e.g., current red-zone congested stations) using modern web design principles, dark mode, and dynamic charts.")

doc.add_page_break()

# Page 4
doc.add_heading('4. Environment Setup', level=1)

doc.add_heading('4.1 Backend Setup', level=2)
doc.add_paragraph("The backend environment was initialized using Python 3.10+ to ensure high performance and modern language features.")
doc.add_paragraph("Virtual Environment: A virtual environment (.venv) was created to isolate dependencies.", style='List Bullet')
doc.add_paragraph("Framework: FastAPI was chosen for its async capabilities and automatic OpenAPI documentation generation.", style='List Bullet')
doc.add_paragraph("Dependencies Installed: fastapi, uvicorn (Server); sqlalchemy, asyncpg (ORM & Database connection); pandas, numpy, scikit-learn (Data processing); pyjwt, passlib (Authentication).", style='List Bullet')
doc.add_paragraph("Directory Structure: Structured into routes, models, schemas, services, and core configurations.", style='List Bullet')

doc.add_heading('4.2 Frontend Setup', level=2)
doc.add_paragraph("The frontend was bootstrapped using Vite and React for a lightning-fast development experience.")
doc.add_paragraph("Initialization: npm create vite@latest frontend -- --template react-ts", style='List Bullet')
doc.add_paragraph("Styling: Vanilla CSS with a predefined design system focusing on glassmorphism and modern UI elements.", style='List Bullet')
doc.add_paragraph("State Management: Context API and React Query for fetching and caching backend data.", style='List Bullet')
doc.add_paragraph("Routing: react-router-dom implemented for navigating between the Dashboard, Station Management, and Settings.", style='List Bullet')

doc.add_heading('4.3 Authentication and Role-based Access', level=2)
doc.add_paragraph("Security is paramount in an operational management system.")
doc.add_paragraph("Implemented JWT (JSON Web Tokens) authentication on the backend.", style='List Bullet')
doc.add_paragraph("Created login and registration endpoints.", style='List Bullet')
doc.add_paragraph("Developed a Role-Based Access Control (RBAC) middleware to restrict access. For instance, only 'Admins' can modify train schedules, while 'Operators' can view dashboards and acknowledge alerts.", style='List Bullet')
doc.add_paragraph("On the frontend, protected routes were configured to redirect unauthenticated users to the login screen.", style='List Bullet')

doc.add_page_break()

# Page 5
doc.add_heading('5. Dataset Collection Strategy', level=1)

doc.add_paragraph("To power the AI and Analytics components of MetroFlow, comprehensive datasets were required. Since real-time production data from metro authorities is often restricted, a rigorous data synthesis and collection process was undertaken.")

doc.add_heading('5.1 Data Sources & Types', level=2)
doc.add_paragraph("We compiled data representing a realistic metro network, focusing on the following core areas:")
doc.add_paragraph("Ticketing Data: Entry and exit logs, tap-in/tap-out timestamps, fare zones.", style='List Bullet')
doc.add_paragraph("Ridership Data: Historical daily, weekly, and monthly passenger counts per station and line.", style='List Bullet')
doc.add_paragraph("Train Schedules: Planned timetables vs. actual arrival/departure times, highlighting delays.", style='List Bullet')
doc.add_paragraph("Occupancy Data: Train car weight sensors or camera estimations providing occupancy percentages.", style='List Bullet')
doc.add_paragraph("Sensor & GPS Data: Location tracking of trains and platform density metrics.", style='List Bullet')

doc.add_heading('5.2 Data Preparation & Cleaning', level=2)
doc.add_paragraph("Raw Data Ingestion: Datasets were stored in the datasets/raw directory in CSV and JSON formats.", style='List Bullet')
doc.add_paragraph("Handling Missing Values: Imputation techniques were used to fill gaps in sensor data, using forward-fill for continuous telemetry and mean-imputation for historical counts.", style='List Bullet')
doc.add_paragraph("Feature Engineering: Extracted time-based features (Hour of day, Day of week, Is_Weekend, Is_Holiday) which are crucial for predicting crowd surges.", style='List Bullet')
doc.add_paragraph("Normalization: Scaled occupancy and density metrics to standard ranges (0-100%) to ensure uniform processing by the visualization engine and future AI models.", style='List Bullet')

doc.add_paragraph("The cleaned and processed datasets were subsequently saved to the datasets/processed directory, ready for Exploratory Data Analysis (EDA) and model training.")

doc.add_page_break()

# Page 6
doc.add_heading('6. Exploratory Data Analysis (EDA) - Ridership & Occupancy', level=1)
doc.add_paragraph("Before building prediction models and the dashboard UI, it was critical to understand the underlying patterns in the passenger data. We conducted an extensive Exploratory Data Analysis (EDA) process.")

doc.add_heading('6.1 Ridership Trends', level=2)
doc.add_paragraph("Analyzing historical ridership helps identify peak hours, seasonal variations, and overall network load. We generated time-series plots to visualize the daily passenger volume across different metro lines. The analysis revealed distinct bimodal distributions corresponding to morning and evening rush hours.")

doc.add_paragraph("Key Performance Indicator (KPI) Dashboard:").bold = True
doc.add_paragraph("The KPI visualizations provided a snapshot of overall network health, total daily riders, and average dwell times.")
try:
    doc.add_picture(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\datasets\charts\ridership\kpi_dashboard.png", width=Inches(6.0))
except Exception as e:
    doc.add_paragraph(f"[Image not found: kpi_dashboard.png]")

doc.add_paragraph("Ridership Year-over-Year Growth:").bold = True
doc.add_paragraph("By comparing historical data, we observed the growth trajectory of the metro system, aiding in long-term capacity planning.")
try:
    doc.add_picture(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\datasets\charts\ridership\yoy_growth.png", width=Inches(6.0))
except Exception as e:
    pass

doc.add_heading('6.2 Train Occupancy Analysis', level=2)
doc.add_paragraph("Train occupancy is a critical metric for crowd management. We analyzed the distribution of occupancy rates across different routes.")

doc.add_paragraph("Occupancy by Line Distribution:").bold = True
doc.add_paragraph("This analysis helped pinpoint which specific lines suffer from chronic overcrowding and which are underutilized.")
try:
    doc.add_picture(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\datasets\charts\occupancy\line_dist.png", width=Inches(6.0))
except Exception:
    pass

doc.add_paragraph("Occupancy Distance Correlation:").bold = True
doc.add_paragraph("We examined if longer routes correlated with higher or lower average occupancies, assisting in route optimization.")
try:
    doc.add_picture(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\datasets\charts\occupancy\distance.png", width=Inches(6.0))
except Exception:
    pass

doc.add_page_break()

# Page 7
doc.add_heading('7. Exploratory Data Analysis (EDA) - Cross-Analysis & Conclusion', level=1)

doc.add_heading('7.1 Cross-Analysis & Congestion Tracking', level=2)
doc.add_paragraph("To build a holistic view of the metro system's performance, we performed cross-analysis between different datasets, such as linking ticketing entries with train delays and platform sensor data.")

doc.add_paragraph("Cross-Dataset Correlation:").bold = True
doc.add_paragraph("By joining ticketing and occupancy datasets, we validated that surges in station entries strongly predict near-future train car congestion in those specific zones.")
try:
    doc.add_picture(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\datasets\charts\cross_analysis\join_coverage.png", width=Inches(6.0))
except Exception:
    pass

doc.add_paragraph("Crowd Management Distributions:").bold = True
doc.add_paragraph("We analyzed the statistical distributions of crowd density scores. This helped in defining the exact numerical thresholds for 'Normal', 'Warning', and 'Critical' congestion levels in our alerting system.")
try:
    doc.add_picture(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\datasets\charts\cross_analysis\crowd_mgmt_distributions.png", width=Inches(6.0))
except Exception:
    pass

doc.add_heading('7.2 Milestone 1 Conclusion', level=2)
doc.add_paragraph("The completion of Milestone 1 establishes a robust foundation for the MetroFlow platform.")
doc.add_paragraph("The project initialization phase successfully defined the scope, architecture, and database schemas required for a high-performance transportation analytics tool.", style='List Bullet')
doc.add_paragraph("The frontend and backend setup yielded a functional, secure, and modern development environment with authentication and basic routing in place.", style='List Bullet')
doc.add_paragraph("The dataset collection and EDA provided invaluable insights into passenger behavior, validating our assumptions and providing the clean data necessary for the upcoming AI/ML modeling phases.", style='List Bullet')
doc.add_paragraph("Moving into Milestone 2, the focus will shift towards training the AI-based crowd prediction models and integrating real-time web sockets to bring the analytics dashboard to life with live telemetry data.")

doc.save(r"d:\Playground\Eat_Sleep_Code_Repeat\Projects\Platform-for-Metro-Crowd-Management-and-Scheduling\docs\Milestone_1_Report.docx")
print("Document saved successfully as Milestone_1_Report.docx.")
