import subprocess, sys, os

def install_and_generate():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # ─── Helpers ───────────────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_table_borders(table, color='4472C4'):
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for side in ['top', 'left', 'bottom', 'right']:
                    be = OxmlElement(f'w:{side}')
                    be.set(qn('w:val'), 'single')
                    be.set(qn('w:sz'), '4')
                    be.set(qn('w:space'), '0')
                    be.set(qn('w:color'), color)
                    tcBorders.append(be)
                tcPr.append(tcBorders)

    def code_block(doc, code_text, label=None):
        """Add a styled code block paragraph."""
        if label:
            lp = doc.add_paragraph()
            lr = lp.add_run(f"  {label}")
            lr.bold = True
            lr.font.color.rgb = RGBColor(100, 100, 100)
            lr.font.size = Pt(9)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 180, 100)  # green code color
        # dark background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E1E1E')
        pPr.append(shd)
        return p

    def note_box(doc, text, color='FFF3CD', border='FFA500'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"⚠  {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(133, 77, 0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        pPr.append(shd)

    def success_box(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"✅  {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 100, 0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E6F9EE')
        pPr.append(shd)

    # ─── Document ──────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    # ════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MetroFlow")
    r.bold = True; r.font.size = Pt(40); r.font.color.rgb = RGBColor(31, 56, 100)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("AI Metro Crowd Management & Scheduling")
    rs.font.size = Pt(16); rs.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_paragraph()
    b = doc.add_paragraph()
    b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = b.add_run("PROJECT RUN GUIDE  ·  VS CODE SETUP & COMMANDS")
    rb.bold = True; rb.font.size = Pt(14); rb.font.color.rgb = RGBColor(255, 255, 255)
    pPr = b._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1F3864')
    pPr.append(shd)

    doc.add_paragraph()
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.style = 'Table Grid'
    meta = [
        ("Project", "MetroFlow — AI Platform for Metro Crowd Management & Scheduling"),
        ("Milestone", "Milestone 1: Week 1 & 2 — Project Initialization, Design & Core Setup"),
        ("Prepared By", "Satti Sri Mahi"),
        ("Date", "July 8, 2026"),
    ]
    for i, (k, v) in enumerate(meta):
        row = meta_tbl.rows[i]
        row.cells[0].text = ''
        row.cells[1].text = v
        rk = row.cells[0].paragraphs[0].add_run(k)
        rk.bold = True; rk.font.color.rgb = RGBColor(31,56,100)
        set_cell_bg(row.cells[0], 'D9E1F2')
        if i % 2 == 0:
            set_cell_bg(row.cells[1], 'F5F7FF')
    add_table_borders(meta_tbl)

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 0 — PREREQUISITES
    # ════════════════════════════════════════
    doc.add_heading("Prerequisites — Install Before Running", level=1)
    doc.add_paragraph(
        "Before running MetroFlow in VS Code, make sure the following software is installed on your machine:"
    )

    pre_tbl = doc.add_table(rows=6, cols=4)
    pre_tbl.style = 'Table Grid'
    headers = ["Software", "Required Version", "Download Link", "Purpose"]
    hrow = pre_tbl.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = ''
        hr = hrow.cells[i].paragraphs[0].add_run(h)
        hr.bold = True; hr.font.color.rgb = RGBColor(255,255,255); hr.font.size = Pt(11)
        set_cell_bg(hrow.cells[i], '1F3864')

    pre_data = [
        ("Python", "3.9 or higher", "python.org/downloads", "Backend API runtime"),
        ("pip", "Latest (bundled with Python)", "Included with Python", "Python package manager"),
        ("VS Code", "Latest", "code.visualstudio.com", "Code editor & terminal"),
        ("Git", "2.x or higher", "git-scm.com", "Version control (optional)"),
        ("Web Browser", "Chrome / Firefox / Edge", "Pre-installed", "Open frontend dashboard"),
    ]
    for i, row_data in enumerate(pre_data):
        row = pre_tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(10.5)
            if j == 0:
                run.bold = True; run.font.color.rgb = RGBColor(31,56,100)
            if i % 2 == 0:
                set_cell_bg(row.cells[j], 'EEF2FF')
    add_table_borders(pre_tbl)

    doc.add_paragraph()
    doc.add_heading("Recommended VS Code Extensions", level=2)
    exts = [
        "Python (by Microsoft) — for Python syntax highlighting and IntelliSense",
        "Pylance — for advanced Python type checking",
        "REST Client (by Huachao Mao) — for testing API endpoints inside VS Code",
        "Live Server — for serving frontend HTML files directly in the browser",
        "SQLite Viewer — to visually inspect metroflow.db database tables",
    ]
    for e in exts:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 1 — PROJECT STRUCTURE
    # ════════════════════════════════════════
    doc.add_heading("1. Project Folder Structure", level=1)
    doc.add_paragraph("After opening the project in VS Code, you will see this folder structure:")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(
        "Platform-for-Metro-Crowd-Management-and-Scheduling/\n"
        "│\n"
        "├── backend/                    ← Python FastAPI server\n"
        "│   ├── main.py                 ← Main API application\n"
        "│   ├── init_db.py              ← Database setup & data migration\n"
        "│   ├── run_eda.py              ← EDA analytics & chart generation\n"
        "│   ├── generate_docx.py        ← Report generation script\n"
        "│   ├── preprocessing.py        ← Data cleaning utilities\n"
        "│   ├── requirements.txt        ← All Python dependencies\n"
        "│   ├── metroflow.db            ← SQLite database (auto-created)\n"
        "│   └── data/                   ← Raw CSV datasets (10 files)\n"
        "│       ├── passenger_flow.csv\n"
        "│       ├── train_occupancy.csv\n"
        "│       ├── train_schedule.csv\n"
        "│       ├── delay_logs.csv\n"
        "│       └── sensor_readings.csv (+ 5 more)\n"
        "│\n"
        "├── frontend/                   ← Browser dashboard\n"
        "│   ├── index.html              ← Main dashboard UI\n"
        "│   └── app.js                  ← JavaScript logic & API calls\n"
        "│\n"
        "└── Milestone_1_Report.docx     ← Project report"
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1E1E1E')
    pPr.append(shd)
    run.font.color.rgb = RGBColor(220, 220, 170)

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 2 — OPEN IN VS CODE
    # ════════════════════════════════════════
    doc.add_heading("2. Opening the Project in VS Code", level=1)

    doc.add_heading("Method A — Using VS Code GUI", level=2)
    steps_a = [
        "Open VS Code application on your laptop",
        "Click File → Open Folder",
        "Navigate to: Downloads → Platform-for-Metro-Crowd-Management-and-Scheduling",
        "Click Open",
        "VS Code will load the full project in the Explorer panel on the left",
        "Open the integrated terminal: View → Terminal  (or press Ctrl + ` )",
    ]
    for i, s in enumerate(steps_a):
        doc.add_paragraph(f"Step {i+1}: {s}", style='List Number')

    doc.add_paragraph()
    doc.add_heading("Method B — Using Terminal Command", level=2)
    doc.add_paragraph("If VS Code is in your system PATH, open any terminal and run:")
    code_block(doc, "code /Users/YourName/Downloads/Platform-for-Metro-Crowd-Management-and-Scheduling")
    note_box(doc, "Replace 'YourName' with your actual macOS/Windows username.")

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 3 — ALL RUN COMMANDS
    # ════════════════════════════════════════
    doc.add_heading("3. Complete Run Commands — Step by Step", level=1)
    doc.add_paragraph(
        "Follow every step below in the VS Code integrated terminal. "
        "Open two separate terminals (Terminal 1 for backend, Terminal 2 for frontend)."
    )

    # ── STEP 1 ──
    doc.add_heading("STEP 1 — Navigate to the Backend Folder", level=2)
    doc.add_paragraph("In Terminal 1, go inside the backend directory:")
    code_block(doc, "cd /Users/mahi/Downloads/Platform-for-Metro-Crowd-Management-and-Scheduling/backend", "macOS/Linux:")
    code_block(doc, "cd C:\\Users\\YourName\\Downloads\\Platform-for-Metro-Crowd-Management-and-Scheduling\\backend", "Windows:")

    # ── STEP 2 ──
    doc.add_paragraph()
    doc.add_heading("STEP 2 — Create a Python Virtual Environment", level=2)
    doc.add_paragraph("Creating a virtual environment keeps project dependencies isolated:")
    code_block(doc, "python3 -m venv venv", "macOS/Linux:")
    code_block(doc, "python -m venv venv", "Windows:")
    success_box(doc, "This creates a 'venv/' folder inside the backend directory.")

    # ── STEP 3 ──
    doc.add_paragraph()
    doc.add_heading("STEP 3 — Activate the Virtual Environment", level=2)
    code_block(doc, "source venv/bin/activate", "macOS/Linux:")
    code_block(doc, "venv\\Scripts\\activate", "Windows:")
    note_box(doc, "After activation, your terminal prompt will show (venv) at the beginning. This confirms the environment is active.")

    # ── STEP 4 ──
    doc.add_paragraph()
    doc.add_heading("STEP 4 — Install All Required Dependencies", level=2)
    doc.add_paragraph("Install every Python library needed by the project from requirements.txt:")
    code_block(doc, "pip install -r requirements.txt")
    doc.add_paragraph()
    doc.add_paragraph("Also install the additional packages used by the project:")
    code_block(doc, "pip install fastapi uvicorn python-multipart python-docx")
    success_box(doc, "All packages installed. This includes FastAPI, Uvicorn, Pandas, Matplotlib, Seaborn, and SQLite interfaces.")

    # ── STEP 5 ──
    doc.add_paragraph()
    doc.add_heading("STEP 5 — Initialize the Database (Run Once)", level=2)
    doc.add_paragraph(
        "This step creates the SQLite database (metroflow.db), sets up all 10 tables, "
        "and imports 81,000+ rows of data from the CSV files in the data/ folder. "
        "Run this only ONCE — or whenever you want to reset the database:"
    )
    code_block(doc, "python3 init_db.py", "macOS/Linux:")
    code_block(doc, "python init_db.py", "Windows:")

    doc.add_paragraph()
    doc.add_paragraph("Expected output after running:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(
        "  ✅ Database created: metroflow.db\n"
        "  ✅ Table: passenger_flow — 12,000 rows loaded\n"
        "  ✅ Table: train_occupancy — 8,500 rows loaded\n"
        "  ✅ Table: train_schedule — 6,200 rows loaded\n"
        "  ✅ Table: delay_logs — 4,100 rows loaded\n"
        "  ✅ Table: sensor_readings — 50,000 rows loaded\n"
        "  ✅ Users seeded: admin, operator\n"
        "  ✅ Database initialization complete!"
    )
    r.font.name = 'Courier New'; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0, 200, 100)
    pPr2 = p._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), '1E1E1E')
    pPr2.append(shd2)

    # ── STEP 6 ──
    doc.add_paragraph()
    doc.add_heading("STEP 6 — Generate EDA Charts & Analysis Report (Optional)", level=2)
    doc.add_paragraph(
        "Run this script to generate the 5 analytical charts in backend/eda_plots/ "
        "and create the EDA summary markdown report. This is optional but recommended:"
    )
    code_block(doc, "python3 run_eda.py", "macOS/Linux:")
    code_block(doc, "python run_eda.py", "Windows:")
    success_box(doc, "5 chart PNG files saved in backend/eda_plots/. Report saved as backend/eda_report.md.")

    doc.add_page_break()

    # ── STEP 7 ──
    doc.add_heading("STEP 7 — Start the Backend API Server  ⬅ MAIN COMMAND", level=2)

    p_imp = doc.add_paragraph()
    r_imp = p_imp.add_run("  🚀  This is the most important command. It starts the MetroFlow server.")
    r_imp.bold = True; r_imp.font.size = Pt(11); r_imp.font.color.rgb = RGBColor(0, 70, 160)
    pPr_imp = p_imp._p.get_or_add_pPr()
    shd_imp = OxmlElement('w:shd')
    shd_imp.set(qn('w:val'), 'clear'); shd_imp.set(qn('w:color'), 'auto'); shd_imp.set(qn('w:fill'), 'D9E1F2')
    pPr_imp.append(shd_imp)

    doc.add_paragraph()
    code_block(doc, "python3 -m uvicorn main:app --reload --port 8000", "macOS/Linux:")
    code_block(doc, "python -m uvicorn main:app --reload --port 8000", "Windows:")

    doc.add_paragraph()
    doc.add_paragraph("Expected output after running:")
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.5)
    r3 = p3.add_run(
        "  INFO:     Will watch for changes in these directories: ['.']\n"
        "  INFO:     Uvicorn running on http://127.0.0.1:8000 (Press CTRL+C to quit)\n"
        "  INFO:     Started reloader process using StatReload\n"
        "  INFO:     Started server process\n"
        "  INFO:     Waiting for application startup.\n"
        "  INFO:     Application startup complete."
    )
    r3.font.name = 'Courier New'; r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(150, 220, 255)
    pPr3 = p3._p.get_or_add_pPr()
    shd3 = OxmlElement('w:shd')
    shd3.set(qn('w:val'), 'clear'); shd3.set(qn('w:color'), 'auto'); shd3.set(qn('w:fill'), '1E1E1E')
    pPr3.append(shd3)

    doc.add_paragraph()
    note_box(doc, "Keep this terminal OPEN while using the app. Do NOT press Ctrl+C — that stops the server.")
    note_box(doc, "The --reload flag means the server auto-restarts when you save code changes.")

    # ── STEP 8 ──
    doc.add_paragraph()
    doc.add_heading("STEP 8 — Start the Frontend Dashboard", level=2)
    doc.add_paragraph("Open a NEW second terminal in VS Code (click the + icon in the terminal panel) and run:")
    code_block(doc, "cd /Users/mahi/Downloads/Platform-for-Metro-Crowd-Management-and-Scheduling/frontend", "Navigate to frontend (macOS):")
    code_block(doc, "cd C:\\Users\\YourName\\Downloads\\Platform-for-Metro-Crowd-Management-and-Scheduling\\frontend", "Navigate to frontend (Windows):")
    doc.add_paragraph()
    code_block(doc, "python3 -m http.server 8080", "Start server (macOS/Linux):")
    code_block(doc, "python -m http.server 8080", "Start server (Windows):")
    success_box(doc, "Frontend server started at: http://localhost:8080")

    # ── STEP 9 ──
    doc.add_paragraph()
    doc.add_heading("STEP 9 — Open Dashboard in Browser", level=2)
    doc.add_paragraph("Open any browser and go to:")
    code_block(doc, "http://localhost:8080")
    doc.add_paragraph("You will see the MetroFlow login page. Use these credentials:")

    cred_tbl = doc.add_table(rows=3, cols=3)
    cred_tbl.style = 'Table Grid'
    ch = cred_tbl.rows[0]
    for i, h in enumerate(["Account Type", "Username", "Password"]):
        ch.cells[i].text = ''
        r = ch.cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(ch.cells[i], '1F3864')
    cred_data = [
        ("🔐  System Admin", "admin", "password"),
        ("🛠️  Station Operator", "operator", "password"),
    ]
    for i, row_data in enumerate(cred_data):
        row = cred_tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(10.5)
            if j == 0:
                run.bold = True
            if i == 0:
                set_cell_bg(row.cells[j], 'E2EFDA')
            else:
                set_cell_bg(row.cells[j], 'FFF2CC')
    add_table_borders(cred_tbl)

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 4 — API ENDPOINTS REFERENCE
    # ════════════════════════════════════════
    doc.add_heading("4. API Endpoints Reference", level=1)
    doc.add_paragraph(
        "Once the backend is running on port 8000, you can test the following endpoints "
        "directly in your browser or using the VS Code REST Client extension:"
    )

    api_tbl = doc.add_table(rows=8, cols=4)
    api_tbl.style = 'Table Grid'
    ah = api_tbl.rows[0]
    for i, h in enumerate(["Method", "Endpoint URL", "Description", "Test in Browser?"]):
        ah.cells[i].text = ''
        r = ah.cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(ah.cells[i], '1F3864')

    api_data = [
        ("GET", "http://127.0.0.1:8000/", "Health check — confirms server is running", "✅ Yes"),
        ("POST", "http://127.0.0.1:8000/auth/login", "Login with username + password", "❌ Use REST Client"),
        ("GET", "http://127.0.0.1:8000/crowd/summary", "Live passenger flow metrics", "✅ Yes"),
        ("GET", "http://127.0.0.1:8000/crowd/alerts", "Active congestion alerts list", "✅ Yes"),
        ("GET", "http://127.0.0.1:8000/crowd/occupancy", "Top station occupancy data", "✅ Yes"),
        ("POST", "http://127.0.0.1:8000/operator/resolve-alert", "Resolve a congestion alert", "❌ Use REST Client"),
        ("GET", "http://127.0.0.1:8000/admin/users", "List all operator accounts", "✅ Yes"),
    ]
    for i, row_data in enumerate(api_data):
        row = api_tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True
                if val == "GET":
                    run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    run.font.color.rgb = RGBColor(180, 70, 0)
            if i % 2 == 0:
                set_cell_bg(row.cells[j], 'EEF2FF')
    add_table_borders(api_tbl)

    doc.add_paragraph()
    doc.add_heading("Auto-Generated API Documentation", level=2)
    doc.add_paragraph("FastAPI automatically creates interactive API documentation. Open in browser:")
    code_block(doc, "http://127.0.0.1:8000/docs", "Swagger UI (Interactive):")
    code_block(doc, "http://127.0.0.1:8000/redoc", "ReDoc (Clean Reference):")

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 5 — CURL TEST COMMANDS
    # ════════════════════════════════════════
    doc.add_heading("5. Testing Commands (curl / Terminal)", level=1)
    doc.add_paragraph("You can test all API endpoints directly from the VS Code terminal using these curl commands:")

    curl_cmds = [
        ("Health Check", 'curl http://127.0.0.1:8000/'),
        ("Admin Login", 'curl -s -X POST -H "Content-Type: application/json" \\\n     -d \'{"username": "admin", "password": "password"}\' \\\n     http://127.0.0.1:8000/auth/login'),
        ("Operator Login", 'curl -s -X POST -H "Content-Type: application/json" \\\n     -d \'{"username": "operator", "password": "password"}\' \\\n     http://127.0.0.1:8000/auth/login'),
        ("Get Crowd Summary", 'curl http://127.0.0.1:8000/crowd/summary'),
        ("Get Active Alerts", 'curl http://127.0.0.1:8000/crowd/alerts'),
        ("Get Station Occupancy", 'curl http://127.0.0.1:8000/crowd/occupancy'),
        ("Get All Users (Admin)", 'curl http://127.0.0.1:8000/admin/users'),
        ("Resolve an Alert", 'curl -s -X POST -H "Content-Type: application/json" \\\n     -d \'{"alert_id": "ALERT-STN-1", "username": "operator", "notes": "Gates re-routed"}\' \\\n     http://127.0.0.1:8000/operator/resolve-alert'),
    ]

    for label, cmd in curl_cmds:
        doc.add_heading(label, level=3)
        code_block(doc, cmd)
        doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 6 — STOP / RESTART
    # ════════════════════════════════════════
    doc.add_heading("6. Stopping & Restarting the Project", level=1)

    doc.add_heading("To Stop the Backend Server:", level=2)
    doc.add_paragraph("In Terminal 1 (where uvicorn is running), press:")
    code_block(doc, "Ctrl + C")
    success_box(doc, "Server stopped. The database (metroflow.db) is safely preserved.")

    doc.add_heading("To Stop the Frontend Server:", level=2)
    doc.add_paragraph("In Terminal 2 (where http.server is running), press:")
    code_block(doc, "Ctrl + C")

    doc.add_heading("To Restart Everything:", level=2)
    doc.add_paragraph("Run these commands again in order:")
    restart_steps = [
        "cd backend",
        "source venv/bin/activate  (macOS) / venv\\Scripts\\activate  (Windows)",
        "python3 -m uvicorn main:app --reload --port 8000",
        "Open new terminal → cd frontend → python3 -m http.server 8080",
    ]
    for i, step in enumerate(restart_steps):
        doc.add_paragraph(f"{i+1}. {step}", style='List Number')

    doc.add_paragraph()
    doc.add_heading("To Reset the Database (Fresh Start):", level=2)
    note_box(doc, "WARNING: This deletes all existing data and re-imports from CSV files.")
    code_block(doc, "rm metroflow.db && python3 init_db.py", "macOS/Linux:")
    code_block(doc, "del metroflow.db && python init_db.py", "Windows:")

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 7 — TROUBLESHOOTING
    # ════════════════════════════════════════
    doc.add_heading("7. Troubleshooting — Common Errors & Fixes", level=1)

    trouble_tbl = doc.add_table(rows=8, cols=3)
    trouble_tbl.style = 'Table Grid'
    th = trouble_tbl.rows[0]
    for i, h in enumerate(["Error / Problem", "Likely Cause", "Fix"]):
        th.cells[i].text = ''
        r = th.cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(th.cells[i], '1F3864')

    trouble_data = [
        ("ModuleNotFoundError: No module named 'fastapi'", "Virtual environment not activated or pip install not run", "Run: pip install -r requirements.txt  inside activated venv"),
        ("Address already in use (port 8000)", "Another process is using port 8000", "Run: lsof -i :8000 to find PID, then kill <PID>. Or change port to --port 8001"),
        ("CORS error in browser console", "Backend CORS not configured for frontend origin", "Confirm CORSMiddleware is present in main.py with allow_origins=['*']"),
        ("Database not found: metroflow.db", "init_db.py was not run yet", "Run python3 init_db.py from the backend/ folder"),
        ("Dashboard shows no data / blank cards", "Backend server is not running", "Start backend first: python3 -m uvicorn main:app --reload --port 8000"),
        ("Permission denied (venv/bin/activate)", "Script not executable on macOS", "Run: chmod +x venv/bin/activate  then try activating again"),
        ("python3 not found (Windows)", "Python installed as 'python' not 'python3'", "Use 'python' instead of 'python3' in all commands on Windows"),
    ]
    for i, row_data in enumerate(trouble_data):
        row = trouble_tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if j == 0:
                run.bold = True; run.font.color.rgb = RGBColor(180, 0, 0)
                run.font.name = 'Courier New'
            if i % 2 == 0:
                set_cell_bg(row.cells[j], 'FFF0F0' if j == 0 else 'EEF2FF')
    add_table_borders(trouble_tbl)

    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 8 — QUICK REFERENCE CARD
    # ════════════════════════════════════════
    doc.add_heading("8. Quick Reference — All Commands at a Glance", level=1)
    doc.add_paragraph("Cut this page out and keep it handy for quick reference during your demonstration:")

    quick_tbl = doc.add_table(rows=13, cols=3)
    quick_tbl.style = 'Table Grid'
    qh = quick_tbl.rows[0]
    for i, h in enumerate(["#", "Action", "Command"]):
        qh.cells[i].text = ''
        r = qh.cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(qh.cells[i], '1F3864')

    quick_data = [
        ("1", "Go to backend folder", "cd backend"),
        ("2", "Create virtual env", "python3 -m venv venv"),
        ("3", "Activate virtual env (Mac)", "source venv/bin/activate"),
        ("4", "Activate virtual env (Win)", "venv\\Scripts\\activate"),
        ("5", "Install dependencies", "pip install -r requirements.txt"),
        ("6", "Install extra packages", "pip install fastapi uvicorn python-multipart python-docx"),
        ("7", "Initialize database (once)", "python3 init_db.py"),
        ("8", "Generate EDA charts", "python3 run_eda.py"),
        ("9", "Start backend server", "python3 -m uvicorn main:app --reload --port 8000"),
        ("10", "Go to frontend folder", "cd ../frontend"),
        ("11", "Start frontend server", "python3 -m http.server 8080"),
        ("12", "Open dashboard URL", "http://localhost:8080"),
    ]
    for i, (num, action, cmd) in enumerate(quick_data):
        row = quick_tbl.rows[i + 1]
        row.cells[0].text = num
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = action
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10.5)
        row.cells[2].text = ''
        rc = row.cells[2].paragraphs[0].add_run(cmd)
        rc.font.name = 'Courier New'; rc.font.size = Pt(10); rc.font.color.rgb = RGBColor(0, 120, 0)
        if i % 2 == 0:
            set_cell_bg(row.cells[0], 'EEF2FF')
            set_cell_bg(row.cells[1], 'EEF2FF')
            set_cell_bg(row.cells[2], 'F0FFF0')
        else:
            set_cell_bg(row.cells[2], 'F8FFF8')
    add_table_borders(quick_tbl)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Footer ──
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("— MetroFlow Project Run Guide — Milestone 1 —")
    r_end.bold = True; r_end.font.color.rgb = RGBColor(31, 56, 100); r_end.font.size = Pt(13)

    doc.add_paragraph()
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.add_run("Submitted By: Satti Sri Mahi").font.size = Pt(11)

    pd2 = doc.add_paragraph()
    pd2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd2.add_run("Milestone 1: Week 1 & 2 — Project Initialization, Design Process & Core Setup").font.size = Pt(10)

    # ─── Save ───────────────────────────────────────────────────
    output_path = "/Users/mahi/Downloads/MetroFlow_Run_Guide.docx"
    doc.save(output_path)
    print(f"\n✅ Run Guide saved: {output_path}")

if __name__ == "__main__":
    install_and_generate()
