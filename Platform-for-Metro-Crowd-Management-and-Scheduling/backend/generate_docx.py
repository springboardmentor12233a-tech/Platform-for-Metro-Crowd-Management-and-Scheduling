import subprocess
import sys
import os

def install_and_generate():
    print("Checking for python-docx...")
    try:
        import docx
    except ImportError:
        print("Installing python-docx library...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Create document
    doc = Document()

    # Style definitions
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page / Header Info
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("MetroFlow: AI Metro Crowd Management & Scheduling")
    run_title.bold = True
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(47, 84, 150) # Dark Blue

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("Milestone 1: Project Initialization, Design Process & Core Setup")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    # Add space
    doc.add_paragraph()

    # Objectives list
    doc.add_heading("Milestone 1 Objectives", level=1)
    objectives = [
        "Define project objectives and transportation workflows.",
        "Design system architecture and database schema.",
        "Create UI wireframes and operational workflow planning.",
        "Setup frontend and backend environments.",
        "Implement authentication and role-based access system.",
        "Build crowd monitoring dashboard.",
        "Develop congestion tracking features."
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_paragraph()

    # Detailed Process of Achieving the Objectives
    doc.add_heading("Process of Achieving Milestone 1", level=1)

    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "To achieve the objectives of Milestone 1, we executed a structured software engineering lifecycle, "
        "migrating the system from basic file reads and mockup definitions into a production-ready, database-driven application. "
        "Here is the step-by-step process used to complete each objective:"
    )

    # Step 1
    doc.add_heading("Step 1: Transportation Workflows and Objectives Definition", level=2)
    p_step1 = doc.add_paragraph()
    p_step1.add_run(
        "We formalized the core goals of the system to reduce passenger wait times and detect bottlenecks. "
        "Two operational workflows were outlined:\n"
        "• Passenger Boarding: Tracking users as they swipe smart cards, logging gate counts, and using train coach weigh-ins to log density percentage.\n"
        "• Operator Override: Algorithmic alert alerts operators of overcrowding, prompting immediate mitigation actions."
    )

    # Step 2
    doc.add_heading("Step 2: System Architecture & Database Schema", level=2)
    p_step2 = doc.add_paragraph()
    p_step2.add_run(
        "We established a decoupled client-server architecture. All IoT data streams are aggregated in an SQLite database (metroflow.db). "
        "We wrote database models and created tables for users, passenger flow, train occupancy, train schedules, delay logs, and alerts resolution. "
        "An automation script (init_db.py) was written to migrate 81,000+ CSV sensor logs into these tables and establish indexing for fast queries."
    )

    # Step 3
    doc.add_heading("Step 3: UI Wireframing & Screen Layouts", level=2)
    p_step3 = doc.add_paragraph()
    p_step3.add_run(
        "We sketched wireframes for the user screens. The layout was structured into a three-column dashboard. "
        "The left column renders live turnstile entries and exits. The right column displays active warning notices and density summaries. "
        "We mapped operational workflows where operators resolve alerts, which logs overrides in the SQLite database and clears the warnings."
    )

    # Step 4
    doc.add_heading("Step 4: Environments Initialization", level=2)
    p_step4 = doc.add_paragraph()
    p_step4.add_run(
        "We initialized the python environment, installing FastAPI, Uvicorn, and SQL interfaces on the backend, "
        "and served the frontend dashboard (HTML5, TailwindCSS, Vanilla JS) via an HTTP web server, establishing robust APIs."
    )

    # Step 5
    doc.add_heading("Step 5: Role-Based Access Control & Logins", level=2)
    p_step5 = doc.add_paragraph()
    p_step5.add_run(
        "We implemented credentials-based authorization querying the users database. "
        "The UI adjusts based on roles: Operators see the mitigation quick actions panel (dispatch trains, broadcast notices) and can resolve alerts. "
        "System Admins are shown the operations users directory table while operator controls are hidden."
    )

    # Step 6 & 7
    doc.add_heading("Step 6 & 7: Crowd Density Dashboard & Algorithms", level=2)
    p_step6 = doc.add_paragraph()
    p_step6.add_run(
        "We developed a live station density tracker that aggregates entries and exits from turnstiles. "
        "A background engine scans sensor logs and triggers warnings when entry flows exceed 120 passengers/min (Warning) "
        "or 180 passengers/min (Critical). Train occupancy alerts trigger at >=95% coach weights."
    )

    # Step 8
    doc.add_heading("Step 8: Exploratory Data Analysis (EDA)", level=2)
    p_step8 = doc.add_paragraph()
    p_step8.add_run(
        "We ran a Python EDA script (run_eda.py) using Matplotlib and Seaborn, generating 5 plots analyzing hourly traffic trends, "
        "busiest lines, train occupancy profiles, delay incident root causes, and sensor outliers. "
        "A comprehensive report (eda_report.md) was created to summarize these analytical results."
    )

    doc.add_paragraph()

    # Submission footer
    doc.add_heading("Submitted By", level=1)
    p_submit = doc.add_paragraph()
    run_name = p_submit.add_run("Satti Sri Mahi\n")
    run_name.bold = True
    run_name.font.size = Pt(12)
    p_submit.add_run("Milestone 1: Week 1 & 2 — Project Initialization, Design Process & Core Setup\n")
    p_submit.add_run("Date: July 8, 2026")

    # Save report
    output_path = "/Users/mahi/Downloads/Platform-for-Metro-Crowd-Management-and-Scheduling/Milestone_1_Report.docx"
    doc.save(output_path)
    print(f"Word Document saved successfully at: {output_path}")

if __name__ == "__main__":
    install_and_generate()
